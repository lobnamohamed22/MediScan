import io
import re
import time
import os
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
}

print("Fetching share URL...")
r = requests.get(share_url, headers=headers, timeout=20)
urls = re.findall(r'https?://[^\s"\'>]+', r.text)
download_url = None
for url in urls:
    if "download.aspx" in url and "tempauth=" in url:
        download_url = url.replace(r'\u0026', '&')
        break

if not download_url:
    print("Download URL not found.")
    exit(1)

print("Downloading document...")
r_file = requests.get(download_url, headers=headers, timeout=30)
if r_file.status_code != 200:
    print("Download failed.")
    exit(1)

doc = docx.Document(io.BytesIO(r_file.content))

updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    if sec.header.paragraphs:
        p = sec.header.paragraphs[0]
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text

# Save to a bytes buffer
buffer = io.BytesIO()
doc.save(buffer)
file_bytes = buffer.getvalue()
total_size = len(file_bytes)
print(f"Total modified file size: {total_size} bytes")

# Split size (1,000,000 bytes)
chunk_size = 1000000
parts = []
offset = 0
part_idx = 0

# Clean old parts first
for f in os.listdir('.'):
    if f.startswith("Lobna_MediScan.part"):
        try:
            os.remove(f)
        except Exception:
            pass

while offset < total_size:
    chunk = file_bytes[offset:offset+chunk_size]
    part_filename = f"Lobna_MediScan.part{part_idx:02d}"
    with open(part_filename, "wb") as f:
        f.write(chunk)
    print(f"Saved {part_filename} ({len(chunk)} bytes)")
    parts.append(part_filename)
    offset += chunk_size
    part_idx += 1

# Write rebuild_doc.py script
rebuild_script_content = """import os

part_prefix = "Lobna_MediScan.part"
output_filename = "Lobna_Mohamed_Mahmoud_MediScan.docx"

# Find all parts
parts = sorted([f for f in os.listdir('.') if f.startswith(part_prefix)])

if not parts:
    print("No part files found!")
    exit(1)

print(f"Found {len(parts)} parts. Merging...")

with open(output_filename, "wb") as outfile:
    for part in parts:
        print(f"Reading {part}...")
        with open(part, "rb") as infile:
            outfile.write(infile.read())

print(f"Successfully assembled: {output_filename}")
print(f"Size: {os.path.getsize(output_filename)} bytes")
"""

with open("assemble_doc.py", "w") as f:
    f.write(rebuild_script_content)
print("Created assemble_doc.py")

# Check if parts exist
print("Verifying part files exist in current directory:")
for part in parts:
    print(f"  {part}: {os.path.exists(part)}")
