import io
import re
import time
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
}

def get_doc_with_9_sections():
    for attempt in range(10):
        print(f"Attempt {attempt+1}: Fetching share URL...")
        try:
            r = requests.get(share_url, headers=headers, timeout=20)
            if r.status_code != 200:
                print(f"  Share page status: {r.status_code}")
                time.sleep(2)
                continue
            
            urls = re.findall(r'https?://[^\s"\'>]+', r.text)
            download_url = None
            for url in urls:
                if "download.aspx" in url and "tempauth=" in url:
                    download_url = url.replace(r'\u0026', '&')
                    break
            
            if not download_url:
                print("  Download URL not found in HTML.")
                time.sleep(2)
                continue
                
            print("  Downloading document...")
            r_file = requests.get(download_url, headers=headers, timeout=45)
            if r_file.status_code != 200:
                print(f"  Download failed status: {r_file.status_code}")
                time.sleep(2)
                continue
                
            doc = docx.Document(io.BytesIO(r_file.content))
            num_sections = len(doc.sections)
            print(f"  Successfully parsed document. Sections: {num_sections}")
            if num_sections == 9:
                return doc
            else:
                print("  Warning: Document has incorrect number of sections. Retrying...")
                
        except Exception as e:
            print(f"  Error on attempt {attempt+1}: {e}")
        time.sleep(3)
    return None

doc = get_doc_with_9_sections()
if not doc:
    print("Failed to download a valid document after all retries.")
    exit(1)

updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    if sec.header.paragraphs:
        p = sec.header.paragraphs[0]
        old_text = p.text
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text
        print(f"Updated Section {sec_idx} header from {repr(old_text)} to {repr(p.text)}")

# Save to a bytes buffer in memory
buffer = io.BytesIO()
doc.save(buffer)
file_bytes = buffer.getvalue()
print(f"Generated docx size in memory: {len(file_bytes)} bytes")

# Upload to transfer.sh
filename = "Lobna_Mohamed_Mahmoud_MediScan.docx"
upload_url = f"https://transfer.sh/{filename}"

print("Uploading to transfer.sh...")
try:
    r = requests.put(upload_url, data=file_bytes, timeout=120)
    print("Upload status:", r.status_code)
    download_link = r.text.strip()
    print("Download Link:", download_link)
    
    # Verify the download link immediately
    if r.status_code == 200 and download_link:
        print("Verifying the download link...")
        r_verify = requests.get(download_link, stream=True, timeout=20)
        print("Verification status code:", r_verify.status_code)
        content_header = r_verify.raw.read(4)
        print("First 4 bytes of verified download:", content_header)
        if content_header == b'PK\x03\x04':
            print("SUCCESS: The download link is active and valid!")
        else:
            print("WARNING: Downloaded file header does not match a ZIP/DOCX package.")
except Exception as e:
    print("Upload/Verification failed:", e)
