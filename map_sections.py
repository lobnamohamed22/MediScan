import io
import re
import time
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"

headers_options = [
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/114.0'
    }
]

def fetch_document():
    html = None
    for attempt in range(5):
        headers = headers_options[attempt % len(headers_options)]
        print(f"Attempt {attempt+1}: Fetching share URL...")
        try:
            r = requests.get(share_url, headers=headers, timeout=20)
            if r.status_code == 200:
                html = r.text
                print("Successfully fetched share URL!")
                break
            else:
                print(f"Fetch failed with status code {r.status_code}")
        except Exception as e:
            print(f"Fetch raised exception: {e}")
        time.sleep(2 * (attempt + 1))
        
    if not html:
        print("Failed to fetch share page after all attempts.")
        return None
        
    urls = re.findall(r'https?://[^\s"\'>]+', html)
    download_url = None
    for url in urls:
        if "download.aspx" in url and "tempauth=" in url:
            download_url = url.replace(r'\u0026', '&')
            break
            
    if not download_url:
        print("Could not find download URL in HTML page.")
        # Save HTML for debugging
        with open("debug_fail.html", "w", encoding="utf-8") as f:
            f.write(html)
        return None
        
    print("Found download URL. Downloading file...")
    for attempt in range(3):
        try:
            r_file = requests.get(download_url, headers=headers, timeout=30)
            if r_file.status_code == 200:
                print("Download successful!")
                return r_file.content
            else:
                print(f"Download failed with status code {r_file.status_code}")
        except Exception as e:
            print(f"Download raised exception: {e}")
        time.sleep(3)
    return None

content = fetch_document()
if not content:
    print("Could not get document content.")
    exit(1)

doc = docx.Document(io.BytesIO(content))
print(f"Total sections: {len(doc.sections)}")

paragraph_sections = []
current_section_idx = 0
section_paragraphs = {i: [] for i in range(len(doc.sections))}

namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for p_idx, p in enumerate(doc.paragraphs):
    section_paragraphs[current_section_idx].append(p)
    pPr = p._p.get_or_add_pPr()
    sectPr = pPr.find('w:sectPr', namespaces)
    if sectPr is not None:
        print(f"Section break found at P{p_idx} (belongs to Section {current_section_idx})")
        current_section_idx += 1

print(f"Remaining paragraphs belong to Section {current_section_idx}")

for sec_idx in range(len(doc.sections)):
    print(f"\n================ SECTION {sec_idx} ================")
    h_text = ""
    if doc.sections[sec_idx].header.paragraphs:
        h_text = " // ".join([p.text.strip() for p in doc.sections[sec_idx].header.paragraphs if p.text.strip()])
    print(f"Header: {repr(h_text)}")
    printed = 0
    for p in section_paragraphs[sec_idx]:
        txt = p.text.strip()
        if txt:
            print(f"  Body: {repr(txt[:120])}")
            printed += 1
            if printed >= 5:
                break
