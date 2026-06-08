import io
import re
import time
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"

headers_options = [
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/114.0'
    }
]

def fetch_document():
    html = None
    for attempt in range(5):
        headers = headers_options[attempt % len(headers_options)]
        print(f"Attempt {attempt+1}: Fetching share URL...")
        try:
            r = requests.get(share_url, headers=headers, timeout=20)
            if r.status_code == 200:
                html = r.text
                print("Successfully fetched share URL!")
                break
            else:
                print(f"Fetch failed with status code {r.status_code}")
        except Exception as e:
            print(f"Fetch raised exception: {e}")
        time.sleep(2 * (attempt + 1))
        
    if not html:
        print("Failed to fetch share page after all attempts.")
        return None
        
    urls = re.findall(r'https?://[^\s"\'>]+', html)
    download_url = None
    for url in urls:
        if "download.aspx" in url and "tempauth=" in url:
            download_url = url.replace(r'\u0026', '&')
            break
            
    if not download_url:
        print("Could not find download URL in HTML page.")
        return None
        
    print("Found download URL. Downloading file...")
    for attempt in range(3):
        try:
            r_file = requests.get(download_url, headers=headers, timeout=30)
            if r_file.status_code == 200:
                print("Download successful!")
                return r_file.content
            else:
                print(f"Download failed with status code {r_file.status_code}")
        except Exception as e:
            print(f"Download raised exception: {e}")
        time.sleep(3)
    return None

content = fetch_document()
if not content:
    print("Could not download document.")
    exit(1)

doc = docx.Document(io.BytesIO(content))
print(f"Loaded document. Total sections: {len(doc.sections)}")

# Define the updates we want to make.
# Section 2: 'Chapter 3: System Analysis' -> 'Chapter 2: Business Analysis & Modeling'
# Section 4: 'Chapter 5: System Design' -> 'Chapter 4: System Modeling'
# Section 6: 'Chapter 7: Conclusions and Future Work' -> 'Chapter 6: System Implementation & Testing'
# Wait! Let's check Section 8 header: Should we change Section 8 header to 'References' or keep it?
# In map_sections, Section 8 was References, header was 'Chapter 7: Conclusions and Future Work'.
# Let's change Section 8 header to 'References' or keep it as is?
# Wait! The user says "Please ensure that all blue header chapter numbers and titles match the corresponding chapter cover pages throughout the document."
# References is not a chapter, but its cover/title page is titled "References". If it has a header showing "Chapter 7...", it doesn't match its own cover page. Let's change Section 8 header to 'References'! This is extremely professional and matches the logic.
# Wait, let's see. Let's make Section 8 header be 'References'. Wait! Let's also check if there are other sections.
# Let's review the changes:
updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    header = sec.header
    # Let's update the first paragraph's run text if it exists
    if header.paragraphs:
        p = header.paragraphs[0]
        old_text = p.text
        if p.runs:
            # Modify the text of the first run and clear the text of other runs to preserve style
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
            print(f"Updated Section {sec_idx} header from {repr(old_text)} to {repr(p.text)}")
        else:
            p.text = new_text
            print(f"Updated Section {sec_idx} header paragraph text directly to {repr(p.text)}")

# Save to the local file
filename = r"c:\Users\lenovo\Downloads\MediScan (4) (1) (1)\Lobna_Mohamed_Mahmoud_MediScan.docx"
doc.save(filename)
print(f"Saved modified document to {filename}")

# Double check
doc_check = docx.Document(filename)
print("\n--- Verifying Headers in saved file ---")
for i in range(len(doc_check.sections)):
    h_text = ""
    if doc_check.sections[i].header.paragraphs:
        h_text = " // ".join([p.text.strip() for p in doc_check.sections[i].header.paragraphs if p.text.strip()])
    print(f"Section {i} header: {repr(h_text)}")
