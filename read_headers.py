import io
import re
import requests
import docx

# 1. Fetch sharing page to get fresh token and URL
share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

print("Fetching share URL...")
try:
    r = requests.get(share_url, headers=headers, timeout=30)
    html = r.text
    urls = re.findall(r'https?://[^\s"\'>]+', html)
    download_url = None
    for url in urls:
        if "download.aspx" in url and "tempauth=" in url:
            download_url = url.replace(r'\u0026', '&')
            break
except Exception as e:
    print("Failed to fetch share page:", e)
    download_url = None

if not download_url:
    # Use fallback direct URL
    download_url = (
        "https://my.microsoftpersonalcontent.com/personal/64c5fd955c53e9e2/_layouts/15/download.aspx"
        "?UniqueId=4c8daf12-20e8-448c-ab05-3a5a18ed5f5b"
        "&Translate=false"
        "&tempauth=v1e.eyJzaXRlaWQiOiI5Y2U5ZjdmZC05ZWM4LTQyYTEtYjMxYS1mOGFiYmZmZDJkYmIiLCJhdWQiOiIwMDAwMDAwMy0wMDAwLTBmZjEtY2UwMC0wMDAwMDAwMDAwMDAvbXkubWljcm9zb2Z0cGVyc29uYWxjb250ZW50LmNvbUA5MTg8MDQwZC02YzY3LTRjNWItYjExMi0zNmEzMDRiNjZkYWQiLCJleHAiOiIxNzgxMjc3MjE0In0.AWUssMAerFbQ5vCCtRoNjAR0y7SqgeUXifG3RAsWpSKZjN9714sfhdyMffFXm3a975Z-09Bqv84FPBwXSGVzqu2pEdHxEGPNqM45hII22_JA5aPjgwdQe4BpD_y-83rJGNZmFpfmjS70NI2UkQT0hX7ToTpeI3ckIAR8jC1Ef82U2glvFqKHbvggAOavWlMTFU4aGjZxxKfJM_hfdrQ1_Cs2up8YIJxo-s5UcBpfv6gt90WU9CAXaEPJ938xrBpInQOgVIl5f9J1EnHF2tznrRG53I6phEu0mvUeQIIyTlr8Ar8TvBzgIJN7oMRwATBGN7xXyFWXhDjuKjRLAYvH0LJK9uSzO-lIg-agWv1alO1QeiRrTBPhnII876cnHiFXBRr_0MuapFYUXlGNwXuTwDNN9EKfwDRBorXTE8no2xFIyHtyKhGMQHY_-Haq1F-_70o8h7mbCvs93okmLNWoPoqZJXH1V2IQeqRYGtIfT6ufcuCO7IQBo5yme0QOTbMnCrePY-V9t_MX0xOuT6kGMrN3ict5mkhF_mf0MSwVqtdxe6ewUrCcdomKNuOPcqAP.YGrd0VxjLRFWH23wkJ0dIIQNwUfeX86ofs0uNM119sc"
    )

print("Downloading file...")
r_file = requests.get(download_url, headers=headers, timeout=60)
print("Status:", r_file.status_code)

if r_file.status_code == 200:
    doc = docx.Document(io.BytesIO(r_file.content))
    print(f"Total sections: {len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        header = sec.header
        print(f"\nSection {i} header:")
        print("  is_linked_to_previous:", sec.header.is_linked_to_previous)
        for pi, p in enumerate(header.paragraphs):
            print(f"  P{pi}: {repr(p.text)}")
            for ri, run in enumerate(p.runs):
                print(f"    Run {ri}: {repr(run.text)}")
        # Check tables in header (sometimes headers use tables for layout)
        for ti, table in enumerate(header.tables):
            print(f"  Table {ti}:")
            for row_i, row in enumerate(table.rows):
                for col_i, cell in enumerate(row.cells):
                    for pi, p in enumerate(cell.paragraphs):
                        if p.text.strip():
                            print(f"    Cell ({row_i},{col_i}) P{pi}: {repr(p.text)}")
else:
    print("Download failed.")
