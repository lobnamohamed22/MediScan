import io
import re
import time
import os
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
}

def get_doc_with_9_sections():
    for attempt in range(10):
        print(f"Attempt {attempt+1}: Fetching share URL...")
        try:
            r = requests.get(share_url, headers=headers, timeout=20)
            if r.status_code != 200:
                time.sleep(2)
                continue
            
            urls = re.findall(r'https?://[^\s"\'>]+', r.text)
            download_url = None
            for url in urls:
                if "download.aspx" in url and "tempauth=" in url:
                    download_url = url.replace(r'\u0026', '&')
                    break
            
            if not download_url:
                time.sleep(2)
                continue
                
            r_file = requests.get(download_url, headers=headers, timeout=40)
            if r_file.status_code != 200:
                time.sleep(2)
                continue
                
            doc = docx.Document(io.BytesIO(r_file.content))
            if len(doc.sections) == 9:
                return doc
        except Exception as e:
            print(f"Error: {e}")
        time.sleep(3)
    return None

doc = get_doc_with_9_sections()
if not doc:
    print("Failed to download doc.")
    exit(1)

updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    if sec.header.paragraphs:
        p = sec.header.paragraphs[0]
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text

filename = "test_output.bin"
doc.save(filename)
print(f"Saved to {filename}. Exists: {os.path.exists(filename)}")

for i in range(10):
    time.sleep(1)
    print(f"Second {i+1}: exists = {os.path.exists(filename)}")
