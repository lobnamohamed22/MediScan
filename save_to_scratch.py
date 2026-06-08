import io
import re
import time
import os
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"

headers_options = [
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/114.0'
    }
]

def fetch_document():
    html = None
    for attempt in range(5):
        headers = headers_options[attempt % len(headers_options)]
        print(f"Attempt {attempt+1}: Fetching share URL...")
        try:
            r = requests.get(share_url, headers=headers, timeout=20)
            if r.status_code == 200:
                html = r.text
                break
        except Exception as e:
            print(f"Fetch exception: {e}")
        time.sleep(2)
        
    if not html:
        return None
        
    urls = re.findall(r'https?://[^\s"\'>]+', html)
    download_url = None
    for url in urls:
        if "download.aspx" in url and "tempauth=" in url:
            download_url = url.replace(r'\u0026', '&')
            break
            
    if not download_url:
        return None
        
    for attempt in range(3):
        try:
            r_file = requests.get(download_url, headers=headers, timeout=30)
            if r_file.status_code == 200:
                return r_file.content
        except Exception as e:
            print(f"Download exception: {e}")
        time.sleep(3)
    return None

content = fetch_document()
if not content:
    print("Failed to download document.")
    exit(1)

doc = docx.Document(io.BytesIO(content))
print(f"Total sections: {len(doc.sections)}")

updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

# Check if we have enough sections
if len(doc.sections) <= max(updates.keys()):
    print(f"Error: Document has only {len(doc.sections)} sections, but we need at least {max(updates.keys())+1} sections.")
    exit(1)

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    if sec.header.paragraphs:
        p = sec.header.paragraphs[0]
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text
        print(f"Updated Section {sec_idx} header to: {p.text}")

scratch_dir = r"C:\Users\lenovo\.gemini\antigravity\brain\4063795c-ef98-479b-9bb7-7d50dfe97afa\scratch"
if not os.path.exists(scratch_dir):
    os.makedirs(scratch_dir)

filename = os.path.join(scratch_dir, "Lobna_Mohamed_Mahmoud_MediScan.docx")
doc.save(filename)
print(f"Saved modified document to: {filename}")

print("Sleeping 5 seconds to check persistence...")
time.sleep(5)
print("File exists after 5 seconds:", os.path.exists(filename))
if os.path.exists(filename):
    print("File size:", os.path.getsize(filename))
