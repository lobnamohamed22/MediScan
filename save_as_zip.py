import io
import re
import time
import os
import requests
import docx

share_url = "https://1drv.ms/w/c/64c5fd955c53e9e2/IQASr41M6CCMRKsFOloY7V9bAVVZNhERVVNU9HdcKBpqVzM?e=8HjDbY"
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
}

print("Fetching share URL...")
r = requests.get(share_url, headers=headers, timeout=20)
urls = re.findall(r'https?://[^\s"\'>]+', r.text)
download_url = None
for url in urls:
    if "download.aspx" in url and "tempauth=" in url:
        download_url = url.replace(r'\u0026', '&')
        break

if not download_url:
    print("Download URL not found.")
    exit(1)

print("Downloading document...")
r_file = requests.get(download_url, headers=headers, timeout=30)
if r_file.status_code != 200:
    print("Download failed.")
    exit(1)

doc = docx.Document(io.BytesIO(r_file.content))

updates = {
    2: "Chapter 2: Business Analysis & Modeling",
    4: "Chapter 4: System Modeling",
    6: "Chapter 6: System Implementation & Testing",
    8: "References"
}

for sec_idx, new_text in updates.items():
    sec = doc.sections[sec_idx]
    if sec.header.paragraphs:
        p = sec.header.paragraphs[0]
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text

# Save it as .zip
filename = "Lobna_Mohamed_Mahmoud_MediScan.zip"
doc.save(filename)
print(f"Saved modified document as {filename}")
print("Exists:", os.path.exists(filename))
